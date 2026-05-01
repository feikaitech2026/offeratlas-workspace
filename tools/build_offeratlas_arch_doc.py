from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\dev-env\workspace\OfferAtlas")
OUT = ROOT / "docs" / "OfferAtlas技术架构与阿里云部署方案.docx"


BLUE = "4F46E5"
GREEN = "10B981"
INK = "111827"
MUTED = "64748B"
LIGHT_BLUE = "EEF2FF"
LIGHT_GREEN = "ECFDF5"
LIGHT_GRAY = "F8FAFC"
LINE = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15


def add_table(doc, headers, rows, widths=None, header_fill=BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_style(table)
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Cm(width)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF", size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_shading(cells[i], "FFFFFF")
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 14 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), size=8.8, align=align)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else INK)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(INK)
        r2 = p.add_run(text[len(bold_prefix):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for r in runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)
        if not r.font.color.rgb:
            r.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_style(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.name = "Microsoft YaHei"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def set_doc_defaults(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("OfferAtlas 技术架构与阿里云部署方案")
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)


def build():
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("OfferAtlas 技术架构与阿里云部署方案")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run("面向留学机构 SaaS 平台的本地仿生产环境与阿里云生产部署规划")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        doc,
        "方案结论",
        "采用 B/C 前端分离、Core API 统一业务后端、ScholarGraph 院校库服务独立、Crawler 数据采集器独立的架构。本地通过 Docker Compose 模拟生产依赖，生产环境部署在阿里云 ACK，并使用 RDS PostgreSQL、Redis、OSS、OpenSearch、SLS、ARMS、ACR、ALB/WAF 等云产品。",
        LIGHT_BLUE,
    )

    add_heading(doc, "1. 项目与服务边界", 1)
    add_table(
        doc,
        ["项目", "定位", "核心职责"],
        [
            ["OfferAtlas Pro", "B 端前端", "顾问、文书、机构管理员使用；学生管理、申请流程、文书、院校匹配、数据分析。"],
            ["OfferAtlas Student", "C 端前端", "学生和家长使用；选校、申请进度、文书确认、消息、费用预算、主题切换。"],
            ["OfferAtlas Core API", "统一业务后端", "租户、用户、学生、申请、文书、材料、消息、权限、套餐、通知、审计。"],
            ["OfferAtlas ScholarGraph", "院校库服务", "院校/专业库、搜索、匹配数据、排名、申请要求、费用、申请截止日期。"],
            ["OfferAtlas Crawler", "数据采集器", "院校官网与公开数据采集、清洗、去重、入库、索引更新、采集日志。"],
        ],
        widths=[4, 3.2, 10],
    )

    add_heading(doc, "2. 总体架构", 1)
    add_para(doc, "架构原则：B/C 端前端必须分离；B/C 端业务后端第一阶段共用 Core API；院校数据、搜索与爬虫从主业务中拆出，避免数据采集任务影响核心业务稳定性。")
    add_bullets(
        doc,
        [
            "B 端用户访问 OfferAtlas Pro，再调用 Core API 完成机构内部业务。",
            "C 端用户访问 OfferAtlas Student，再调用 Core API 完成学生门户业务。",
            "Core API 调用 ScholarGraph 获取院校、专业、搜索、匹配与推荐解释。",
            "Crawler 将原始数据写入采集库和暂存表，经清洗审核后进入 ScholarGraph 正式数据表与搜索索引。",
            "文件统一走对象存储，数据库只保存元数据、权限和版本信息。",
        ],
    )

    add_table(
        doc,
        ["层级", "组件", "说明"],
        [
            ["入口层", "CDN / WAF / ALB / Ingress", "处理公网访问、HTTPS、防护、负载均衡与路由。"],
            ["前端层", "Pro Web / Student Web", "两个独立前端项目，分别适配 B 端管理后台与 C 端学生门户。"],
            ["业务层", "Core API", "统一业务模型和权限控制，避免 B/C 后端重复建模。"],
            ["数据中台", "ScholarGraph", "院校知识库、搜索、匹配数据和推荐解释能力。"],
            ["采集层", "Crawler Scheduler / Worker", "异步采集和清洗，不暴露公网。"],
            ["数据层", "PostgreSQL / Redis / OSS / Search", "关系数据、缓存队列、文件、搜索索引。"],
            ["观测层", "SLS / ARMS / Prometheus", "日志、链路追踪、指标与告警。"],
        ],
        widths=[2.5, 5, 10],
    )

    add_heading(doc, "3. 本地 Docker Compose 仿生产环境", 1)
    add_para(doc, "本地环境尽量与生产依赖保持一致，避免开发时使用轻量替代导致上线后行为不一致。建议将所有基础设施配置放入 infra/docker-compose。")
    add_table(
        doc,
        ["服务", "本地组件", "生产对应"],
        [
            ["反向代理", "Nginx", "ALB + Ingress"],
            ["B 端前端", "offeratlas-pro", "ACK Pod"],
            ["C 端前端", "offeratlas-student", "ACK Pod"],
            ["业务后端", "offeratlas-core-api", "ACK Pod"],
            ["院校服务", "offeratlas-scholargraph", "ACK Pod"],
            ["爬虫", "crawler-scheduler / crawler-worker", "ACK Job/CronJob/Worker Pod"],
            ["数据库", "PostgreSQL", "RDS PostgreSQL"],
            ["缓存/队列", "Redis", "Tair/Redis"],
            ["对象存储", "MinIO", "OSS"],
            ["搜索", "Meilisearch 或 Elasticsearch", "OpenSearch/Elasticsearch"],
            ["邮件测试", "Mailpit", "邮件推送服务或 SMTP"],
            ["监控", "Prometheus + Grafana", "ARMS + SLS + 云监控"],
        ],
        widths=[3, 6, 7],
    )
    add_para(doc, "本地域名建议：pro.offeratlas.localhost、student.offeratlas.localhost、api.offeratlas.localhost、scholar.offeratlas.localhost、minio.offeratlas.localhost。")

    add_heading(doc, "4. 数据库与存储设计", 1)
    add_para(doc, "生产环境建议使用同一个 RDS PostgreSQL 实例，但拆分多个 database，保证边界清晰，同时便于未来单独迁出 ScholarGraph 或 Crawler。")
    add_table(
        doc,
        ["Database", "用途", "说明"],
        [
            ["offeratlas_core", "核心业务库", "租户、用户、学生、申请、文书、材料、消息、套餐、审计。"],
            ["offeratlas_scholar", "院校正式数据", "院校、专业、排名、费用、申请要求、截止日期、匹配特征。"],
            ["offeratlas_crawler", "采集与暂存库", "原始页面、采集文件、暂存院校/专业/排名、采集任务日志。"],
        ],
        widths=[4, 4.5, 9],
    )
    add_bullets(
        doc,
        [
            "OSS 存储学生材料、文书附件、Offer PDF、导出报告和爬虫原始文件。",
            "Redis 用于会话缓存、验证码、短期锁、任务队列和热点数据缓存。",
            "搜索引擎维护院校、专业、申请要求、标签、全文检索和筛选索引。",
            "所有文件访问通过签名 URL 或后端授权代理，禁止公开桶直连敏感材料。",
        ],
    )

    add_heading(doc, "5. 推荐技术栈", 1)
    add_table(
        doc,
        ["领域", "推荐技术", "说明"],
        [
            ["前端", "Next.js / React / TypeScript", "适合 B/C 两端独立工程，支持 SSR、路由和工程化。"],
            ["样式", "Tailwind CSS + shadcn/ui 或 Ant Design", "B 端偏管理系统，C 端通过 Design Token 支持白/黑主题。"],
            ["状态与请求", "TanStack Query + Zustand", "服务端状态和局部 UI 状态分离。"],
            ["后端", "NestJS / TypeScript", "模块化结构清晰，适合长期演进为服务拆分。"],
            ["ORM", "Prisma 或 TypeORM", "配合 PostgreSQL migration 管理。"],
            ["富文本", "TipTap 或 Lexical", "支持文书编辑、版本、批注和 AI 建议应用。"],
            ["搜索", "Meilisearch 起步，OpenSearch/ES 生产", "通过 Search Adapter 屏蔽实现差异。"],
            ["队列", "BullMQ 或 RabbitMQ", "用于爬虫、文档导出、通知、AI 批处理。"],
            ["对象存储", "MinIO / OSS", "本地 MinIO，生产 OSS，使用统一 Storage Adapter。"],
        ],
        widths=[3.2, 5.2, 9],
    )

    add_heading(doc, "6. 阿里云生产采购清单", 1)
    add_para(doc, "以下规格以“正式上线初期、支持真实客户使用”为基准。后续可按租户数量、学生数量、文件量和搜索索引规模扩容。")
    add_table(
        doc,
        ["类别", "阿里云产品", "起步规格建议"],
        [
            ["容器运行", "ACK 托管版", "1 个集群，2 个 worker 节点"],
            ["计算节点", "ECS", "2 台，4 vCPU / 8GB"],
            ["镜像仓库", "ACR", "个人版或企业基础版"],
            ["数据库", "RDS PostgreSQL", "2 vCPU / 4GB 或 4 vCPU / 8GB"],
            ["数据库存储", "RDS ESSD", "200GB 起"],
            ["缓存", "Tair/Redis", "1GB-2GB 起"],
            ["对象存储", "OSS 标准存储", "500GB 起"],
            ["搜索", "OpenSearch / Elasticsearch", "100GB 起；预算敏感时可先自建小规格搜索服务"],
            ["负载均衡", "ALB", "1 个公网 ALB"],
            ["安全", "WAF + SSL 证书", "公网域名防护与 HTTPS"],
            ["日志", "SLS", "应用日志、访问日志、爬虫日志，保留 30 天起"],
            ["监控", "ARMS + 云监控", "APM、链路追踪、告警"],
            ["网络", "VPC / NAT 网关", "私网隔离，服务统一出网"],
            ["加速", "CDN", "静态资源与 OSS 文件下载加速"],
        ],
        widths=[3, 5, 9],
    )

    add_heading(doc, "7. 容量规划", 1)
    add_table(
        doc,
        ["资源", "内测/初期", "商业化初期", "增长期"],
        [
            ["OSS 文件存储", "500GB", "1TB-2TB", "5TB 起"],
            ["RDS 存储", "200GB", "500GB", "1TB 起"],
            ["Redis", "1GB-2GB", "4GB", "8GB+"],
            ["搜索索引", "50GB-100GB", "200GB-500GB", "1TB+"],
            ["ACK 节点", "2 台 4C8G", "3 台 4C16G", "5 台以上或弹性扩容"],
        ],
        widths=[4, 4, 4, 5],
    )
    add_para(doc, "学生文件容量估算：轻量按 50MB/学生，常规按 100MB/学生，重度按 300MB/学生。建议第一期按 OSS 500GB、RDS 200GB、搜索 100GB 起步。")

    add_heading(doc, "8. 阿里云部署拓扑", 1)
    add_bullets(
        doc,
        [
            "公网域名：pro.offeratlas.com、student.offeratlas.com、api.offeratlas.com。",
            "内部服务：ScholarGraph、Crawler、PostgreSQL、Redis、OpenSearch 均不直接暴露公网。",
            "前端静态资源可由 ACK 服务或对象存储 + CDN 承载，API 统一走 ALB/Ingress。",
            "Crawler 以 CronJob、Worker Pod 或独立任务队列方式运行，避免和 Core API 争抢资源。",
            "RDS、Redis、OpenSearch 仅允许 ACK 节点安全组访问。",
        ],
    )

    add_heading(doc, "9. CI/CD 与环境划分", 1)
    add_para(doc, "建议三套环境：dev 使用本地 Docker Compose，staging 使用阿里云测试环境，prod 使用阿里云生产环境。")
    add_table(
        doc,
        ["阶段", "动作", "结果"],
        [
            ["代码提交", "GitHub/GitLab 触发 CI", "运行 lint、类型检查、单元测试。"],
            ["镜像构建", "Docker Build", "为 Pro、Student、Core API、ScholarGraph、Crawler 构建镜像。"],
            ["镜像发布", "推送到 ACR", "按版本号、commit sha、环境 tag 管理。"],
            ["部署", "Helm/K8s Manifest", "部署到 ACK，滚动更新。"],
            ["验证", "健康检查与冒烟测试", "失败自动回滚或阻断发布。"],
        ],
        widths=[3, 5.5, 9],
    )

    add_heading(doc, "10. 迁移阿里云步骤", 1)
    add_bullets(
        doc,
        [
            "第一步：本地 Docker Compose 跑通所有服务和基础依赖。",
            "第二步：统一环境变量，确保本地与生产只换配置、不改代码。",
            "第三步：创建 VPC、交换机、安全组、ACK、ACR、RDS、Redis、OSS、ALB、SLS。",
            "第四步：创建 RDS database：offeratlas_core、offeratlas_scholar、offeratlas_crawler。",
            "第五步：执行数据库 migration 和基础种子数据。",
            "第六步：推送所有服务镜像到 ACR，并部署 staging。",
            "第七步：将 MinIO 替换为 OSS，将本地搜索替换为 OpenSearch/ES 或生产搜索服务。",
            "第八步：完成登录、上传、搜索、文书、申请流转、爬虫任务的端到端测试。",
            "第九步：开启 HTTPS、WAF、备份、日志、监控、告警。",
            "第十步：部署 production，进行灰度发布和回滚演练。",
        ],
    )

    add_heading(doc, "11. 安全与合规", 1)
    add_callout(
        doc,
        "敏感数据范围",
        "系统会处理护照号、手机号、邮箱、成绩单、签证材料、家庭财务材料、Offer、合同、付款信息等敏感数据，必须从第一天纳入权限、审计、加密和访问控制。",
        LIGHT_GREEN,
    )
    add_bullets(
        doc,
        [
            "所有业务表带 tenant_id，并在服务层和数据库查询层做租户隔离。",
            "B 端采用 RBAC：机构管理员、顾问、文书老师、申请老师、运营等角色分权。",
            "C 端学生和家长只能访问授权范围内的申请、材料、文书和消息。",
            "OSS 使用私有桶，文件访问通过签名 URL，敏感文件不公开直链。",
            "关键操作写入 audit_log，包括登录、导出、下载、删除、权限变更、申请状态变更。",
            "生产数据库禁止公网访问，RDS、Redis、OpenSearch 仅允许私网安全组访问。",
            "Crawler 需要遵守目标站 robots、限速和来源记录，采集数据保留来源 URL 与时间。",
            "备份策略：RDS 自动备份 7-30 天，OSS 开启版本控制或生命周期策略，重要文件可做跨区域备份。",
        ],
    )

    add_heading(doc, "12. 分阶段实施路线", 1)
    add_table(
        doc,
        ["阶段", "目标", "交付内容"],
        [
            ["阶段 1：工程地基", "跑通本地仿生产环境", "Docker Compose、Core API、数据库 migration、MinIO/Redis/Search、统一鉴权。"],
            ["阶段 2：核心业务", "完成留学申请闭环", "学生管理、院校库、院校匹配、申请流程、文书、文件、通知。"],
            ["阶段 3：AI 与数据", "形成产品差异化", "AI 文书评分、AI 选校、风险预警、ScholarGraph 数据清洗和搜索优化。"],
            ["阶段 4：生产部署", "迁移阿里云并可正式商用", "ACK、RDS、OSS、Redis、SLS、ARMS、WAF、CI/CD、备份告警。"],
            ["阶段 5：规模化", "按业务增长拆服务", "AI Service、Doc Service、Notification Service、Billing Service 逐步独立。"],
        ],
        widths=[3.5, 4.5, 9],
    )

    add_heading(doc, "13. 当前建议优先事项", 1)
    add_bullets(
        doc,
        [
            "补建 OfferAtlas Core API 目录，作为 B/C 两端共用业务后端。",
            "建立 infra 目录，沉淀 Docker Compose、Nginx、K8s、数据库初始化脚本。",
            "确定 PostgreSQL database 边界和第一版 migration。",
            "抽象 Storage Adapter 和 Search Adapter，保证 MinIO/OSS、Meilisearch/OpenSearch 可替换。",
            "先完成统一登录、租户隔离、RBAC、审计日志，再进入大规模业务页面开发。",
        ],
    )

    add_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
